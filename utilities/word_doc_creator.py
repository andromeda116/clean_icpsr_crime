import pandas as pd
import docx


def create_icpsr_value_counts_word_doc(ip_file, var_list, op_file):
    df = pd.read_csv(ip_file)

    counts_doc = docx.Document()

    for var in var_list:
        counts_doc.add_heading(f'{var} value counts')
        counts_doc_para = counts_doc.add_paragraph(f'Frequencies of different possible vaues of {var}')

        # print([df[f'{col}'].value_counts() for col in list(df) if col.endswith(f'{var}')])

        counts = [df[f'{col}'].value_counts() for col in list(df) if col.endswith(f'{var}')]

#        counts_doc_para.add_run(counts, style = 'List Number')

        counts_df = pd.DataFrame(counts).reset_index()
        print(type(counts_df))
        print(counts_df.head())

        # for cnt in counts:
        #     #counts_doc_para_data = counts_doc.add_paragraph(counts)
        #     print(type(cnt))
        #     counts_doc.add_paragraph(cnt, style='List Number')
        # counts_doc.add_page_break()

    counts_doc.save(op_file)